# _*_ coding:utf-8 _**
# Author:vic yang
# Date: 2017-9
# AutoOfficer
# Version:2.2.0

import win32com
import CmdFormat
from win32com.client import Dispatch, constants
import shutil
import gc
import re
import copy
import os
from multiprocessing import Pool,cpu_count,cpu_count,freeze_support
from ProgressBar import ProgressBar
import time
import socket
import multiProcessPackage
import ReRange,sys
import docx

VERSION = '1.0.1'
COMPANY = "中冶三勘院"
CDMF = CmdFormat.CmdFormat("自动Ofiice v"+VERSION+" 特供赟哥")
ISOTIMEFORMAT='%Y-%m-%d %X'

def log(x):
	""" recording the status information"""
	if x==None:
		return
	with open('操作结果.txt', 'a+') as f:
		f.write(str(x)+'\n')

def calculate_fails():
	"""calcualte the fail opreration number"""
	failNumber = 0
	with open('操作结果.txt', 'r') as f:
		for line in f:
			# regrard "Log Time " as a start flag to calculate fails number
			# because the log file contains many other opreation history logs
			if "Log Time" in line:
				failNumber = 0
			if bool(re.search(r'\d', line)) and ("X" in line): #"X" is a fail flag
				failNumber +=1
	return failNumber

class easyWord(object):
	"""A class for opreating word file"""
	def __init__(self, FileName, bUseDocx):
		self.PagesCount = 0
		self.bUseDocx = bUseDocx
		self.WordAPP = win32com.client.DispatchEx('Word.Application')
		self.WordAPP.Visible = 0
		self.WordAPP.DisplayAlerts = 0
		self.FileName = FileName
		self.Doc = self.WordAPP.Documents.Open(FileName)
		self.PagesCount=self.WordAPP.ActiveWindow.ActivePane.Pages.Count
		# print(FileName)
		tempFame = self.FileName
		if self.bUseDocx :
			if not FileName.endswith('x'):
				tempFName = os.path.dirname(self.FileName)+'\\'+'t.docx'
				self.Doc.SaveAs(tempFName,16)
			self.f = open(tempFName, 'rb')
			self.Docx=docx.Document(self.f)
			self.personNumber = re.sub("\D", "", self.Docx.tables[1].cell(0,7).text)
			self.leader = self.Docx.tables[0].cell(0,10).text
	def pages_count(self):
		return self.PagesCount
	def set_cell(self,R,C,Value,TableIndex=0,FontSize=-1):
		#Range是一个非常重要的概念，可以设置字体，行间距，文本！！
		self.Doc.Tables[TableIndex].Rows[R].Cells[C].Range.Text = Value
		if FontSize!=-1:
			self.Doc.Tables[TableIndex].Rows[R].Cells[C].Range.Font.Size = FontSize
	def get_cell(self,R,C,TableIndex=0):
		return self.Doc.Tables[TableIndex].Rows[R].Cells[C].Range.Text
	def get_person_number(self):
		return int(self.personNumber)
		# return self.Docx.tables[1].cell(0,2).text
		# return self.Doc.Tables[1].Rows[0].Cells[2].Range.Text
	def get_leader(self):
		return self.leader
		# return self.Docx.tables[0].cell(0,3).text
		#return self.Doc.Tables[0].Rows[0].Cells[3].Range.Text
	def close(self):
		if self.bUseDocx:
			self.f.close()
		self.Doc.Save()
		self.Doc.Close()
		self.WordAPP.Quit()
		del self.WordAPP
		if os.path.exists(os.path.dirname(self.FileName)+'\\'+'t.docx'):
			os.remove(os.path.dirname(self.FileName)+'\\'+'t.docx')


class easyExcel(object):
	"""A class for opreating word file"""
	def __init__(self, FileName):
		self.ExcelApp = win32com.client.DispatchEx('Excel.Application')
		self.ExcelApp.Visible=False
		self.ExcelApp.DisplayAlerts = False
		if FileName:
			self.FileName = FileName
			self.Xls = self.ExcelApp.Workbooks.Open(self.FileName)
		else:
			log("文件 "+FileName+" 没找到！")
			raise IOError("文件 "+FileName+" 没找到！")
			return
	def close(self):
		self.Xls.Close(SaveChanges=0)
		self.ExcelApp.Quit()
		del self.ExcelApp
	def pages_count(self):
		pages = 0
		for x in range(1,self.Xls.Worksheets.Count+1):
			Activesheet = self.Xls.Worksheets(x)
			pages = pages + Activesheet.PageSetup.Pages.Count
			#pages = pages + (Activesheet.VPageBreaks.Count)*(Activesheet.HPageBreaks.Count)
		return pages
	# reading areacode and time from "地区代码及时间表"
	def read_areacode_time(self):
		AreaTimeAdict = {}
		for x in range(1,self.Xls.Worksheets.Count+1):
			nValidRows = 0
			Activesheet = self.Xls.Worksheets(x)
			CDMF.print_blue_text("提取 "+Activesheet.Name +" 信息...",end='')
			log("提取 "+Activesheet.Name +" 信息...")
			# UsedRange 从1开始
			nTempCode = Activesheet.Cells(1,1).Value
			nRows = Activesheet.UsedRange.Rows.Count
			nColumns = Activesheet.UsedRange.Columns.Count
			for i in range(2,nRows+1):
				tempList=[]
				for j in range(1,nColumns+1):
					string = Activesheet.Cells(i,j).Value
					if string!=None:
						if j==1:
							key = re.sub("\D", "", string)
							strRe = (re.split(r'([县镇乡村])',string)[2:6])
							strRe.append(str(nTempCode)) #加入乡镇编号，用于填充“软卷皮封面.doc”中的分类号
							tempList.append(strRe)
						elif j<=8 and j>1:
							# translate 2017.1.2  to  20170102
							time = string.split('.')
							if len(time)!=3:
								continue
							if len(time[1])==1:
								time[1] = '0'+time[1]
							if len(time[2])==1:
								time[2] = '0'+time[2]
							tempList.append("".join(time))
						else:
							tempList.append(string) #村的编号，用于填充“软卷皮封面.doc”中的分类号
				if not tempList:
					continue
				nValidRows = nValidRows + 1
				#这是一个疑点,为什么要加一个.copy()，没有弄清楚还
				AreaTimeAdict[key]=tempList.copy()
				# print((''.join(AreaTimeAdict[key])[0])+"保定华北勘查")
				# os.system('pause')
				tempList.clear()
			CDMF.print_blue_text("成功, 共有 "+str(nValidRows)+" 个村庄.")
			log("成功, 共有 "+str(nValidRows)+" 个村庄.")
		CDMF.print_blue_text("有效村庄共有 "+str(len(AreaTimeAdict))+" 个.")
		log("有效村庄共有 "+str(len(AreaTimeAdict))+" 个.")
		return AreaTimeAdict

# 核心函数--------------------------------------------------------------------------------------------
# 实例方法，类方法不能被pickle(序列化)，apply_asyn()需要序列化数据
def tasks(fileOrDir,RootPath,AreaTimeAdict,bRegenerate,bWithTime,bCopyFengmian,nTotal,ProcessOrder):
	# 得到户主名字
	start=time.clock()
	HuZhu = fileOrDir.split('_')[1]
	# 获取户主所在村庄的编号，前12位
	HuZhuVillageCode = (fileOrDir.split('_')[0])[0:12]   #个人所在的村的Code
	HuzhuPersonalCode= (fileOrDir.split('_')[0])[-3:]  #个人的Code，可以填在“软卷皮封面.doc”中的案卷号中
	PersionalDir = os.path.join(RootPath,fileOrDir)
	FileOrDirInPersionalDir_list = os.listdir(os.path.join(RootPath,fileOrDir))
	filesCount=0
	#到这行，农业局确权档案卷内目录.doc肯定存在，无需判断

	try:
		if bRegenerate:
			if os.path.exists(PersionalDir+"\\农业局确权档案卷内目录.doc"):
				os.remove(PersionalDir+"\\农业局确权档案卷内目录.doc")
			shutil.copyfile(RootPath+"\\农业局确权档案卷内目录.doc", PersionalDir+"\\农业局确权档案卷内目录.doc")
		else:
			if os.path.exists(PersionalDir+"\\农业局确权档案卷内目录.doc"):
				pass
			else:
				shutil.copyfile(RootPath+"\\农业局确权档案卷内目录.doc", PersionalDir+"\\农业局确权档案卷内目录.doc")
	except Exception as e:
		raise("WRONG!!!")

	#到这行，软卷皮封面.doc肯定存在或者self.CopyFengmian为false
	try:
		if bCopyFengmian:
			if os.path.exists(PersionalDir +"\\软卷皮封面.doc"):
				os.remove(PersionalDir +"\\软卷皮封面.doc")
			shutil.copyfile(RootPath +"\\软卷皮封面.doc", PersionalDir +"\\软卷皮封面.doc")
	except Exception as e:
		raise("WRONG!!!")
	# 获取所需数据
	Pages_adict = {}
	bWordOpen=False
	bExcelOpen = False
	try:
		#对于一个特定的村民文件夹
		for file in os.listdir(os.path.join(RootPath,fileOrDir)):
			if not os.path.isdir(file) and file.startswith(('1','2','3','4','5','6','7','8','9','0')):
				filesCount = filesCount + 1
				(filepath,tempfilename) = os.path.split(file)
				(filename,extension) = os.path.splitext(tempfilename)
				#print('filename:  '+filename)
				fullName = PersionalDir+'\\'+file
				# Word = easyWord(fullName,False,True)
				if extension==r".doc" :
					if '公示结果归户表' in file:
						Word = easyWord(fullName,True)
					else:
						Word = easyWord(fullName,False)
				elif extension==r'.docx':
					Word = easyWord(fullName,True)
				else:
					continue
				bWordOpen = True
	# 			print('----------------------')
				Pages_adict[filename] = Word.pages_count()
	# # 			# if '登记簿' in file:
	# # 			# 	CunZhang=Word.get_leader()[:-2] #去掉最后两个字符：一个是BEL,一个是换行
	# # 			# 	PersonNumber = Word.get_person_number()
				if '公示结果归户表' in file:
					PersonNumber = Word.get_person_number()
					CunZhang=Word.get_leader() #去掉最后两个字符：一个是BEL,一个是换行
					# print('村长：'+str(CunZhang))
				Word.close()
			# continue
		nTemp = 0
		bFirst = True
		for x in Pages_adict.keys():
			if '声明书' in x:
				nTemp = nTemp + 1
			if '承包合同' in x:
				nTemp = nTemp + 1
			if '地块调查表' in x and bFirst:
				nTemp = nTemp + 1
				bFirst = False
			if '公示结果归户表' in x:
				nTemp = nTemp + 1
		if nTemp != 4: raise IOError
	except Exception as e:
		if bWordOpen:
			Word.close()
		if bExcelOpen:
			Excel.close()
		finish=time.clock()
		infoString = ' '+str(ProcessOrder)+'/'+str(nTotal)+'     ' +fileOrDir + "       操作失败" +"          "+str(finish-start)[0:6]
		CDMF.print_red_text(infoString)
		if os.path.exists(PersionalDir +'\\'+"农业局确权档案卷内目录.doc"):
			os.remove(PersionalDir +'\\'+"农业局确权档案卷内目录.doc")
		log(infoString + "     X")
		return
	# 更新“农业局确权档案卷内目录.doc”
	try:
		Word = easyWord(PersionalDir+'\\'+"农业局确权档案卷内目录.doc",False)

		# 填写目录的第1顺序号
		nTotalPages = 1
		Word.set_cell(1,2,HuZhu) # "责任者"
		if HuZhuVillageCode not in AreaTimeAdict.keys():
			CDMF.print_red_text('在'+'<地区代码及时间表>中没有找到待处理村民所在村编码，请检查<地区代码及时间表>')
			exit(1)
		if bWithTime:
			Word.set_cell(1,4,(AreaTimeAdict[HuZhuVillageCode])[1])  # "日期"
		Word.set_cell(1,5,nTotalPages)     # "页号"

		# 填写目录的第2顺序号
		for x in Pages_adict.keys():
			if '声明书' in x:
				nTotalPages = Pages_adict[x]+nTotalPages
				break
		Word.set_cell(2,2,''.join((AreaTimeAdict[HuZhuVillageCode])[0][0:4])+COMPANY)
		if bWithTime:
			Word.set_cell(2,4,(AreaTimeAdict[HuZhuVillageCode])[2])  # "日期"
		Word.set_cell(2,5,nTotalPages)
		# 填写目录的第3顺序号
		bDjb = False
		for x in Pages_adict.keys():
			if '承包方调查表' in x:
				nTotalPages = Pages_adict[x]+nTotalPages
				bDjb = True
				break
		if not bDjb:
			nTotalPages = nTotalPages + 1 #如果没有承包方调查表，默认承包方调查表为1页
		Word.set_cell(3,2,''.join((AreaTimeAdict[HuZhuVillageCode])[0][0:4])+COMPANY)
		if bWithTime:
			Word.set_cell(3,4,(AreaTimeAdict[HuZhuVillageCode])[3])  # "日期"
		Word.set_cell(3,5,nTotalPages)
		# 填写目录的第4顺序号
		for x in Pages_adict.keys():
			if '地块调查表' in x:
				nTotalPages = Pages_adict[x]+nTotalPages
		Word.set_cell(4,2,''.join((AreaTimeAdict[HuZhuVillageCode])[0][0:4])+COMPANY)
		if bWithTime:
			Word.set_cell(4,4,(AreaTimeAdict[HuZhuVillageCode])[4])  # "日期"
		Word.set_cell(4,5,nTotalPages)
		# 填写目录的第5顺序号
		for x in Pages_adict.keys():
			if '公示结果归户表' in x:
				nTotalPages = Pages_adict[x]+nTotalPages
				#承包方推荐证明
		Word.set_cell(5,2,''.join((AreaTimeAdict[HuZhuVillageCode])[0][0:4])+HuZhu)
		if bWithTime:
			Word.set_cell(5,4,(AreaTimeAdict[HuZhuVillageCode])[5])  # "日期"
		Word.set_cell(5,5,nTotalPages)
		nTotalPages = nTotalPages+1
		# 填写目录的第6顺序号
		# Word.set_cell(6,2,CunZhang+HuZhu)
		Word.set_cell(6,2,''.join((AreaTimeAdict[HuZhuVillageCode])[0][0:4])+HuZhu)
		if bWithTime:
			Word.set_cell(6,4,(AreaTimeAdict[HuZhuVillageCode])[6])  # "日期"
		Word.set_cell(6,5,nTotalPages)
		for x in Pages_adict.keys():
			if '承包合同' in x:
				nTotalPages = Pages_adict[x]+nTotalPages
		# 填写目录的第7顺序号
		#户主户口本及身份证复印件
		Word.set_cell(7,2,HuZhu)
		if bWithTime:
			Word.set_cell(7,4,(AreaTimeAdict[HuZhuVillageCode])[7])  # "日期"
		Word.set_cell(7,5,str(nTotalPages)+"-"+str(nTotalPages+PersonNumber))
		nTotalPages = nTotalPages + PersonNumber
		Pages_adict.clear()
		Word.close()
	except Exception as e:
		infoString = "更新 "+PersionalDir+'\\'+"\"农业局确权档案卷内目录.doc\" 发生错误."
		CDMF.print_red_text(infoString)
		log(infoString)
		return
	# 更新软卷皮封面.doc”
	try:
		if bCopyFengmian:
			Word = easyWord(PersionalDir+'\\'+"软卷皮封面.doc",False)
			mxxx = re.split(r'([镇乡村土])', Word.get_cell(2,0))  # r'([镇乡村土])'加括号保留分隔符
			mxxx[0] = '\r'+(AreaTimeAdict[HuZhuVillageCode])[0][0]
			mxxx[1] = (AreaTimeAdict[HuZhuVillageCode])[0][1]
			mxxx[2] = (AreaTimeAdict[HuZhuVillageCode])[0][2]
			mxxx[4] = HuZhu
			Word.set_cell(2,0,''.join(mxxx),FontSize=18)
			mxxx = re.split(r'([自年(月至)])', Word.get_cell(3,0))
			if mxxx[-1]!='月':
				mxxx.pop()
			if mxxx[0]!='自':
				del mxxx[0]
			if len(mxxx) < 10:
				infoString = "需要检查软卷皮封面模板中“自X年X月至X年X月”"
				log(infoString)
				return
			mxxx[1] = ''.join(list((AreaTimeAdict[HuZhuVillageCode])[7])[0:4])
			if list((AreaTimeAdict[HuZhuVillageCode])[7])[4]=='0':
				mxxx[3] = ''.join(list((AreaTimeAdict[HuZhuVillageCode])[7])[5:6])
			else:
				mxxx[3] = ''.join(list((AreaTimeAdict[HuZhuVillageCode])[7])[4:6])
			mxxx[7] = ''.join(list((AreaTimeAdict[HuZhuVillageCode])[1])[0:4])
			if list((AreaTimeAdict[HuZhuVillageCode])[1])[4]=='0':
				mxxx[9] = ''.join(list((AreaTimeAdict[HuZhuVillageCode])[1])[5:6])
			else:
				mxxx[9] = ''.join(list((AreaTimeAdict[HuZhuVillageCode])[1])[4:6])
			Word.set_cell(3,0,''.join(mxxx),FontSize=18)
			mxxx = re.split(r'([共件页])', Word.get_cell(4,0))
			if mxxx[-1]!='页':
				mxxx.pop()
			if mxxx[0]!='本卷':
				del mxxx[0]
			if len(mxxx) < 6:
				infoString = "需要检查软卷皮封面模板中“本卷共X件X页”"
				log(infoString)
				return
			mxxx[2] = '   7   '
			mxxx[4] = '   '+str(nTotalPages)+'   '
			Word.set_cell(4,0,''.join(mxxx),FontSize=18)
			#设置表2的全宗号
			Word.set_cell(1,0,'53',TableIndex=1,FontSize=12)#小四
			#设置表2的分类号
			strTemp = 'TQ0202'+(AreaTimeAdict[HuZhuVillageCode])[0][4]+str(int((AreaTimeAdict[HuZhuVillageCode])[8]))
			Word.set_cell(1,1,strTemp,TableIndex=1,FontSize=12)#小四
			#设置表2的案卷号
			#Word.set_cell(1,2,str(HuzhuPersonalCode),TableIndex=1,FontSize=12)#小四
			Word.set_cell(1,2,str(ProcessOrder),TableIndex=1,FontSize=12)#小四
			Word.close()
	except Exception as e:
		infoString = "更新 "+PersionalDir+'\\'+"\"软卷皮封面.doc\" 发生错误."
		log(infoString)
		return
	finish=time.clock()
	infoString = ' '+str(ProcessOrder)+'/'+str(nTotal)+'     ' +fileOrDir + "       操作成功" +"          "+str(finish-start)[0:6]
	print(infoString)
	log(infoString)
	#return
# ---------------------------------------------------------------------------------------------------

# def check(RootPath,bRegenerate,bCopyFengmian,bWithTime,AreaTimeAdict,Processes,Status):

class Job(object):
	"""docstring for Job"""
	def __init__(self, RootPath,bRegenerate,bCopyFengmian,bWithTime,AreaTimeAdict,Processes,Status):
			self.bRegenerate = bRegenerate
			self.bCopyFengmian = bRegenerate
			self.bWithTime=bWithTime
			self.AreaTimeAdict = AreaTimeAdict
			self.Status = Status
			self.RootPath = RootPath+"\\"
			self.Processes = Processes
			#一些提示
			CDMF.print_blue_text("扫描待统计村民资料...,")
			self.nNumFile = 0;
			self.nNumNoContent = 0;
			self.filesOrDirsInRoot = os.listdir(self.RootPath)
			for fileOrDir in self.filesOrDirsInRoot:
				if os.path.isdir(os.path.join(self.RootPath,fileOrDir)) and fileOrDir.startswith(('1','2','3','4','5','6','7','8','9','0')):
					self.nNumFile = self.nNumFile + 1
					if not os.path.exists(self.RootPath + fileOrDir + "\\" +"农业局确权档案卷内目录.doc"): #扫描没有目录的
						self.nNumNoContent  = self.nNumNoContent + 1

			CDMF.print_blue_text("扫描完毕！共有 "+str(self.nNumFile) + " 户的资料,",end='')
			log("扫描完毕！共有 "+str(self.nNumFile) + " 户的资料,")
			if self.bRegenerate:
				self.nTotal = self.nNumFile
			else:
				self.nTotal = self.nNumNoContent
			CDMF.print_blue_text("需要统计的有 "+str(self.nTotal) + " 户.")
			log("需要统计的有 "+str(self.nTotal) + " 户.")
			if self.nTotal==0:
				infoString = "没有需要统计的村民."
				CDMF.print_blue_text(infoString)
				quit = input("按任意键退出...")
				self.Status = False
				log(infoString)
				return

			self.Status = True

	def run(self):
		timeString = time.strftime( ISOTIMEFORMAT, time.localtime(time.time()))
		hostName = socket.gethostname()
		log('-----------Log Time: '+timeString+ ' from '+hostName+'  --------------')
		CDMF.print_yellow_text("----------------------------------------------------------------------")
		CDMF.print_yellow_text(" 序号         户主编号与名字                操作状态         耗时(秒)")
		CDMF.print_yellow_text("----------------------------------------------------------------------")
		log("----------------------------------------------------------------------")
		log(" 序号         户主编号与名字                操作状态         耗时(秒)")
		log("----------------------------------------------------------------------")

		#多进程
		try:
			multiP = Pool(self.Processes)
			ProcessOrder = 0
			for fileOrDir in self.filesOrDirsInRoot:
				if os.path.isdir(os.path.join(self.RootPath,fileOrDir)) and fileOrDir.startswith(('1','2','3','4','5','6','7','8','9','0')):
					if not self.bRegenerate and os.path.exists(self.RootPath + fileOrDir + "\\" +"农业局确权档案卷内目录.doc"):
						continue
					ProcessOrder += 1
					multiP.apply_async(tasks, args=(fileOrDir,self.RootPath,self.AreaTimeAdict,self.bRegenerate,self.bWithTime,self.bCopyFengmian,self.nTotal,ProcessOrder,),callback=log)
			multiP.close()
			multiP.join()
			self.Status = True
		except Exception as e:
			infoString = "运行出现错误！"
			CDMF.print_red_text(infoString)
			self.Status = False
			log(infoString)
			return
		if self.Status:
			CDMF.print_yellow_text("----------------------------------------------------------------------")
			failNumber = calculate_fails()
			CDMF.print_yellow_text("共统计 "+ str(self.nTotal) +" 户，成功 "+str(self.nTotal-failNumber)+" 户，失败 "+str(failNumber)+" 户")
			log("----------------------------------------------------------------------")
			log("共统计 "+ str(self.nTotal) +" 户，成功 "+str(self.nTotal-failNumber)+" 户，失败 "+str(failNumber)+" 户")
		else:
			infoString = "运行出现错误！"
			CDMF.print_red_text()
			quit = input("按任意键退出...")
			self.Status = False
			log(infoString)
			return

if __name__ == '__main__':
	freeze_support()

	startTime=time.clock()

	CDMF.print_yellow_text("正在初始化，请稍等......")
	rootpath = os.getcwd()
	ReRange.ReRange(rootpath)
	FilesInRoot = os.listdir(rootpath)
	nCountryNum =0
	for x in FilesInRoot:
		if os.path.isdir(x) and x.endswith('村'):
			nCountryNum += 1
			if os.path.exists(os.path.join(rootpath,x,"农业局确权档案卷内目录.doc")):
				os.remove(os.path.join(rootpath,x,"农业局确权档案卷内目录.doc"))
			shutil.copyfile(os.path.join(rootpath,"农业局确权档案卷内目录.doc"), os.path.join(rootpath,x,"农业局确权档案卷内目录.doc"))
			if os.path.exists(os.path.join(rootpath,x,"软卷皮封面.doc")):
				os.remove(os.path.join(rootpath,x,"软卷皮封面.doc"))
			shutil.copyfile(os.path.join(rootpath,"软卷皮封面.doc"), os.path.join(rootpath,x,"软卷皮封面.doc"))
			# if os.path.exists(os.path.join(rootpath,x,"地区代码及时间表.xlsx")):
			# 	os.remove(os.path.join(rootpath,x,"地区代码及时间表.xlsx"))
			# shutil.copyfile(os.path.join(rootpath,"地区代码及时间表.xlsx"), os.path.join(rootpath,x,"地区代码及时间表.xlsx"))
	CDMF.print_yellow_text("共有 "+str(nCountryNum)+" 村需要处理")

	bRegenerate = True
	bWithTime = False
	bCopyFengmian = True
	AreaTimeAdict= {}
	Status = True
	Processes = [1,]

#	check(rootpath,bRegenerate,bCopyFengmian,bWithTime,AreaTimeAdict,Processes,Status)
#	***********  Check ******************************************************************************
	CDMF.set_cmd_color(CmdFormat.FOREGROUND_RED | CmdFormat.FOREGROUND_GREEN | \
		CmdFormat.FOREGROUND_BLUE | CmdFormat.FOREGROUND_INTENSITY)
	print("\n")
	print("===================   自动Ofiice v"+VERSION+"   =============================")
	print("|                                                                      |")
	print("|      将本程序放在根目录，运行之前请确保根目录下具有                  |")
	CDMF.print_red_text("|      (1) *包含每个村民的个人目录                                     |")
	CDMF.print_red_text("|      (2) *必须具有\"农业局确权档案卷内目录\"文件模板(.doc或.docx)      |")
	print("|      (3) 可以添加\"地区代码及时间表\"文件模板(.xls或.xlsx)             |")
	print("|      (4) 可以添加\"软卷皮封面\"文件模板(.xls或.xlsx)                   |")
	print("|                                                                      |")
	print("========================================================================")
	bWithTime = True
	Status  = True
	RootPath = rootpath + '\\'
	try:
		filesOrDirsInRoot = os.listdir(RootPath)
	except Exception:
		print("No such path: "+RootPath)
	else:
		#print("扫描 "+ self.RootPath)
		if os.path.exists(RootPath+"农业局确权档案卷内目录.doc"):
			pass
		else:
			infoString = "在 "+ RootPath + " 没有找到\"农业局确权档案卷内目录.doc\""
			CDMF.print_red_text(infoString)
			log(infoString)
			CDMF.print_red_text("程序中断，请完善相应资料！")
			Status = False
			quit = input("按任意键退出...")
			sys.exit(0)

		# 询问是否全部重新计算
		while True:
			content = CDMF.print_green_input_text("是否需要新建或重新生成所有目录? 请输入y/Y或者n/N:")
			if content=="y" or content=="Y" or content=="n" or content=="N":
				break
		if content=="y" or content=="Y":
			bRegenerate = True
		else:
			bRegenerate = False

		# 询问是否需要将“软卷皮封面.doc”考入个人目录
		while True:
			content = CDMF.print_green_input_text("是否需要在个人目录自动填充\"软卷皮封面.doc\"? 请输入y/Y或者n/N:")
			if content=="y" or content=="Y" or content=="n" or content=="N":
				break
		if content=="y" or content=="Y":
			bCopyFengmian = True
			if not os.path.exists(RootPath+"软卷皮封面.doc"):
				infoString = "在 "+ RootPath + " 没有找到\"软卷皮封面.doc\""
				log(infoString + "  已忽略此项")
				CDMF.print_red_text(infoString)
				while True:
					content1 = CDMF.print_green_input_text("是否继续? 请输入y/Y或者n/N:")
					if content1=="y" or content1=="Y" or content1=="n" or content=="N":
						break
				if content1=="y" or content1=="Y":
					bCopyFengmian = False
				else:
					CDMF.print_red_text("程序中断，请完善相应资料！")
					quit = input("按任意键退出...")
					Status = False
					log(infoString + "  请完善！")
					sys.exit(0)
		else:
			bCopyFengmian = False

		if bWithTime and "地区代码及时间表.xlsx" in filesOrDirsInRoot:
			CDMF.print_blue_text("正在读取 \"地区代码及时间表.xlsx\"...")
			log("正在读取 \"地区代码及时间表.xlsx\"...")
			Excel = easyExcel(RootPath+"地区代码及时间表.xlsx")
			try:
				AreaTimeAdict = Excel.read_areacode_time()
			except Exception as e:
				infoString = "读取 \"地区代码及时间表.xlsx\" 出错！请检查该文件是否符合模板要求！"
				CDMF.print_red_text(infoString)
				log(infoString)
				bWithTime = False
				CDMF.print_red_text("程序中断，请完善相应资料！")
				quit = input("按任意键退出...")
				Status = False
				sys.exit(0)
			Excel.close()
		elif bWithTime and "地区代码及时间表.xls" in filesOrDirsInRoot:
			CDMF.print_blue_text("正在读取 \"地区代码及时间表.xls\"...")
			log("正在读取 \"地区代码及时间表.xls\"...")
			Excel = easyExcel(RootPath+"地区代码及时间表.xls")
			try:
				AreaTimeAdict = Excel.read_areacode_time()
			except Exception as e:
				infoString = "读取 \"地区代码及时间表.xlsx\" 出错！请检查该文件是否符合模板要求！"
				CDMF.print_red_text(infoString)
				log(infoString)
				bWithTime = False
				CDMF.print_red_text("程序中断，请完善相应资料！")
				quit = input("按任意键退出...")
				Status = False
				sys.exit(0)
			Excel.close()
		elif bWithTime or "地区代码及时间表.xls" not in filesOrDirsInRoot  or "地区代码及时间表.xlsx"not in filesOrDirsInRoot:
			infoString="在 "+RootPath+" 没有找到\"地区代码及时间表.xlsx\"或\"地区代码及时间表.xls\",无法完成目录表中时间自动填充!\""
			CDMF.print_red_text(infoString)
			while True:
				content = CDMF.print_green_input_text("是否继续? 请输入y/Y或者n/N:")
				if content=="y" or content=="Y" or content=="n" or content=="N":
					break
			if content=="y" or content=="Y":
				bWithTime = False
				log(infoString + '  已忽略此项，只在目录中统计页数')
			else:
				CDMF.print_red_text("程序中断，请完善相应资料！")
				log(infoString + '  请完善！')
				quit = input("按任意键退出...")
				Status = False
				sys.exit(0)
		Processes[0] = 1
		while True:
			content1 = CDMF.print_green_input_text("是否启动加速? 请输入y/Y或者n/N:")
			if content1=="y" or content1=="Y" or content1=="n" or content=="N":
				break
		if content1=="y" or content1=="Y":
			CDMF.print_yellow_text("你的CPU核心数为 "+ str(cpu_count()) + ".")
			while True:
				x = CDMF.print_green_input_text("请输入您需要同时算几户？(建议为"+str(cpu_count())+"): ")
				if x.isdigit() and int(x)>=1:
					Processes[0] = int(x)
					break
				else:
					CDMF.print_yellow_text("请输入>=1的整数!")
		else:
			Processes[0] = 1
#		*************************************
	if not AreaTimeAdict:
		CDMF.print_red_text("<地区代码及时间表>读取错误! 请检查该文件!")
		sys.exit(0)
	index = 0
	for x in FilesInRoot:
		if os.path.isdir(os.path.join(rootpath,x)) and x.endswith('村'):
			index += 1
			print("正在处理  " + str(index)+'/'+str(nCountryNum) + "    " + x)
			Jobb = Job(os.path.join(rootpath,x),bRegenerate,bCopyFengmian,bWithTime,AreaTimeAdict,Processes[0],Status)
			if(Jobb.Status):
				Jobb.run()
			del Jobb
			sys.exit(1)

	finishTime = time.clock()
	CDMF.print_yellow_text("总耗时："+str(finishTime-startTime)[0:7]+" 秒")
	CDMF.print_blue_text("任务完成，可查看生成的 “操作结果.txt” 日志.")
	log("总耗时："+str(finishTime-startTime)[0:7]+" 秒")
	quit = input("按任意键退出...")